import streamlit as st
import os
from PIL import Image
import io
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# إعدادات الصفحة
st.set_page_config(page_title="أرشيف وصانع ألبومات الوورد", page_icon="📁", layout="wide")

st.title("📁 ألبوم صور المشاريع وصانع ملفات الوورد الاحترافية")
st.markdown("---")

# المسارات الأساسية للتخزين
BASE_DIR = "photos"
BC_DIR = os.path.join(BASE_DIR, "Bons_de_Commande")
MARCHE_DIR = os.path.join(BASE_DIR, "Marches")

for folder in [BC_DIR, MARCHE_DIR]:
    os.makedirs(folder, exist_ok=True)

# القائمة الجانبية
st.sidebar.header("🔍 تصفح وأرشِف")
choice = st.sidebar.radio("اختر نوع المشروع:", ["سندات الطلب (Bons de Commande)", "الصفقات (Les Marchés)"])
target_dir = BC_DIR if "Bons de Commande" in choice else MARCHE_DIR

# قراءة المشاريع الحالية
projects = [d for d in os.listdir(target_dir) if os.path.isdir(os.path.join(target_dir, d))]

# إمكانية إضافة مشروع جديد ديريكت
new_project = st.sidebar.text_input("➕ إنشاء مشروع جديد (إسم أو رقم):")
if st.sidebar.button("إنشاء الدوسي"):
    if new_project:
        os.makedirs(os.path.join(target_dir, new_project), exist_ok=True)
        st.sidebar.success(f"✅ تم إنشاء دوسي: {new_project}")
        st.rerun()

# تصفح وعرض الألبوم
if projects:
    selected_project = st.sidebar.selectbox("اختر المشروع الحالي:", projects)
    project_path = os.path.join(target_dir, selected_project)
    
    st.header(f"🏗️ المشروع الحالي: {selected_project}")
    
    # 📝 خانات معلومات الـ BC أو المارشي لي غاتبان الفوق د الورقة
    st.markdown("### 📝 معلومات وثيقة الوورد (التقرير)")
    col_info1, col_info2 = st.columns(2)
    with col_info1:
        doc_number = st.text_input("📄 رقم سند الطلب / الصفقة (N° de BC / Marché):", value=selected_project)
    with col_info2:
        doc_subject = st.text_input("🎯 موضوع سند الطلب / الصفقة (Objet):", placeholder="مثال: أشغال التطهير السائل، الإنارة العمومية...")
    
    st.markdown("---")
    
    # رفع صور جديدة للدوسي
    uploaded_files = st.file_uploader("📸 إرفع الصور هنا لحفظها وتوليد ملف الوورد (يمكنك اختيار عدة صور معاً):", accept_multiple_files=True, type=["png", "jpg", "jpeg"])
    
    if uploaded_files:
        for uploaded_file in uploaded_files:
            save_path = os.path.join(project_path, uploaded_file.name)
            with open(save_path, "wb") as f:
                f.write(uploaded_file.getvalue())
        st.success("✅ تم حفظ الصور بنجاح!")
        st.rerun()
    
    # قراءة الصور الموجودة في الدوسي لعرضها وصنع الوورد منها
    valid_extensions = (".png", ".jpg", ".jpeg", ".webp")
    images_in_folder = [img for img in os.listdir(project_path) if img.lower().endswith(valid_extensions)]
    
    if images_in_folder:
        st.write(f"📊 عدد الصور في هذا المشروع: {len(images_in_folder)}")
        
        # 📥 قسم توليد ملف الوورد للطباعة
        st.markdown("### 🖨️ تحميل الألبوم على شكل ملف Word جاهز")
        
        # خيارات التنسيق
        cols_per_page = st.radio("كيف بغيتي تستاف ديال التصاور فالصفحة د الوورد؟", ["صورة واحدة كبيرة في كل صفحة ونقاد ف الوسط دياالها", "جوج صور متناسقين في كل صفحة ونقاد ف الوسط"], index=1)
        
        if st.button("🪄 صاوب ليا ملف الوورد دابا"):
            with st.spinner("جاري إعداد ملف الوورد وترتيب الصور ف الوسط..."):
                doc = Document()
                
                # إعدادات الهوامش
                sections = doc.sections
                for section in sections:
                    section.top_margin = Inches(0.6)
                    section.bottom_margin = Inches(0.6)
                    section.left_margin = Inches(0.6)
                    section.right_margin = Inches(0.6)
                
                # ترويسة معلومات الـ Bon de commande الفوق د الورقة
                header_p = doc.add_paragraph()
                header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                
                run_type = header_p.add_run(f"نوع الإجراء: {choice}\n")
                run_type.font.name = 'Arial'
                run_type.font.size = Pt(12)
                
                run_num = header_p.add_run(f"رقم سند الطلب / الصفقة: {doc_number}\n")
                run_num.font.name = 'Arial'
                run_num.font.size = Pt(14)
                run_num.bold = True
                
                run_obj = header_p.add_run(f"الموضوع: {doc_subject if doc_subject else 'غير محدد'}\n")
                run_obj.font.name = 'Arial'
                run_obj.font.size = Pt(12)
                run_obj.italic = True
                
                # خط فاصل
                p_line = doc.add_paragraph()
                p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_line.add_run("__________________________________________________________________")
                doc.add_paragraph("\n")
                
                # إضافة الصور ف الوورد وتوسيطها بالكامل
                if "صورة واحدة" in cols_per_page:
                    for img_name in images_in_folder:
                        img_full_path = os.path.join(project_path, img_name)
                        try:
                            with Image.open(img_full_path) as test_img:
                                # التصويرة ف السنطر
                                p_img = doc.add_paragraph()
                                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                p_img.add_run().add_picture(img_full_path, width=Inches(6.2))
                                
                                # العنوان ف السنطر ديريكت تحتها
                                p_caption = doc.add_paragraph()
                                p_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                run_cap = p_caption.add_run(f"📸 العنوان: {os.path.splitext(img_name)[0]}")
                                run_cap.font.name = 'Arial'
                                run_cap.font.size = Pt(12)
                                run_cap.bold = True
                                
                                doc.add_page_break()
                        except Exception as e:
                            continue
                else:
                    # صورتين في الصفحة
                    added_count = 0
                    for img_name in images_in_folder:
                        img_full_path = os.path.join(project_path, img_name)
                        try:
                            with Image.open(img_full_path) as test_img:
                                # التصويرة ف السنطر
                                p_img = doc.add_paragraph()
                                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                p_img.add_run().add_picture(img_full_path, width=Inches(4.8))
                                
                                # العنوان ف السنطر تحتها
                                p_caption = doc.add_paragraph()
                                p_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                run_cap = p_caption.add_run(f"📸 العنوان: {os.path.splitext(img_name)[0]}")
                                run_cap.font.name = 'Arial'
                                run_cap.font.size = Pt(11)
                                run_cap.bold = True
                                
                                added_count += 1
                                
                                if added_count % 2 == 0:
                                    doc.add_page_break()
                                else:
                                    doc.add_paragraph("\n")
                        except Exception as e:
                            continue
                
                bio = io.BytesIO()
                doc.save(bio)
                bio.seek(0)
                
                st.download_button(
                    label="📥 اضغط هنا لتحميل ملف الوورد الموسط والمقاد (.docx)",
                    data=bio,
                    file_name=f"Rapport_{doc_number.replace(' ', '_')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
                st.balloons()
        
        st.markdown("---")
        st.markdown("### 👁️ معاينة الصور وإدارتها (تقدر تمسح من هنا):")
        
        # عرض شبكة الصور مع ميزة الحذف تحت كل تصويرة
        cols = st.columns(3)
        for idx, img_name in enumerate(images_in_folder):
            img_path = os.path.join(project_path, img_name)
            try:
                with cols[idx % 3]:
                    st.image(Image.open(img_path), caption=img_name, use_column_width=True)
                    # زر الحذف الخاص بكل صورة
                    if st.button(f"🗑️ مسح {img_name}", key=f"del_{img_name}_{idx}"):
                        os.remove(img_path)
                        st.success(f"تم مسح الصورة: {img_name}")
                        st.rerun()
            except:
                pass
                
    else:
        st.info("ℹ️ هاد الدوسي مازال ما فيه حتا شي تصويرة. إرفع أول تصويرة الفوق!")
else:
    st.info("ℹ️ ابدأ بإنشاء أول دوسي مشروع من القائمة الجانبية.")
